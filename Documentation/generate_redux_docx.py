import os
import sys
import subprocess

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set the background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell in twips."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text, title="NOTE FOR FUTURE SELF", bg_color="F0FDF4", border_color="16A34A"):
    """Create a styled callout box with a colored left border and light background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Set borders: left thick, others none
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')  # 3pt width
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_color)
    tcBorders.append(left)
    
    for side in ['top', 'bottom', 'right']:
        edge = OxmlElement(f'w:{side}')
        edge.set(qn('w:val'), 'nil')
        tcBorders.append(edge)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(22, 163, 74) if border_color == "16A34A" else RGBColor(37, 99, 235)
    
    run_text = p.add_run(text)
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = RGBColor(31, 41, 55)
    
    doc.add_paragraph()  # spacing after table

def add_code_block(doc, code_text):
    """Add a shaded code block table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{side}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '4')
        edge.set(qn('w:color'), 'E2E8F0')
        tcBorders.append(edge)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph()

def create_redux_documentation():
    doc = docx.Document()
    
    # Set standard page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("REDUX TOOLKIT (RTK)\nTHE COMPLETE DEVELOPER GUIDE & NOTES")
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("A Comprehensive Reference Covering Core Concepts, Architecture, Data Flow, & Real-World Integration in the DailyTrack Project")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    # --- SECTION 1: PHASE 1 RESEARCH & EXPLORATION ---
    doc.add_heading("1. Phase 1: Research & Exploration (In My Own Words)", level=1)
    
    doc.add_heading("What is Redux?", level=2)
    doc.add_paragraph(
        "At its core, Redux is a predictable state container for JavaScript applications. When building modern interactive web applications (especially with React), managing 'state'—such as logged-in user details, UI theme, active filters, and data lists—can quickly become overwhelming. Without a centralized system, state is scattered across dozens of isolated components, requiring tedious 'props drilling' (passing props down many parent-child layers).\n\n"
        "Redux solves this by creating a single, centralized 'Store'—a universal source of truth for the entire application. Any component in the app can directly read data from this store and dispatch actions to update it, making state transitions completely predictable and traceable."
    )
    
    doc.add_heading("Why was Redux Toolkit (RTK) introduced & What problems does it solve?", level=2)
    doc.add_paragraph(
        "While traditional Redux was revolutionary, developers frequently complained about three major problems:\n"
        "1. Too Much Boilerplate: Writing standard Redux required creating separate files for action types, action creators, reducers, and store configurations just to perform a simple update.\n"
        "2. Complex Store Setup: Configuring middleware (like redux-thunk for asynchronous requests) and connecting the Redux DevTools required complicated manual boilerplate.\n"
        "3. Manual Immutability: In classic Redux, state could never be modified directly. Developers had to use cumbersome object spreading (e.g., {...state, user: {...state.user, name: 'New'}}) which was prone to accidental state mutations and hard-to-trace bugs.\n\n"
        "Redux Toolkit (RTK) was introduced by the official Redux team as the modern standard to solve all these issues. It provides powerful utility functions like configureStore and createSlice that eliminate boilerplate, bundle Redux DevTools and Thunks automatically, and integrate the 'Immer' library under the hood. Thanks to Immer, we can write direct, intuitive mutable code (like state.goals.push(newGoal)) while RTK automatically generates safe, immutable updates!"
    )
    
    doc.add_heading("Core Building Blocks Defined", level=2)
    
    p = doc.add_paragraph()
    p.add_run("• The Store: ").bold = True
    p.add_run("The central database or repository of your frontend application. It holds the complete global state tree. There is strictly only one store per application, configured using RTK's configureStore().")
    
    p = doc.add_paragraph()
    p.add_run("• A Slice: ").bold = True
    p.add_run("A modular division of the Redux state tree representing a single feature or domain (e.g., studySlice, authSlice, themeSlice). Created using createSlice(), it bundles the initial state, reducer functions, and auto-generated action creators into a single cohesive file.")
    
    p = doc.add_paragraph()
    p.add_run("• Reducers: ").bold = True
    p.add_run("Pure functions that determine how the state changes in response to an action. A reducer takes the current state and an action payload, processes the logic, and calculates the new state.")
    
    p = doc.add_paragraph()
    p.add_run("• Actions: ").bold = True
    p.add_run("Plain JavaScript objects that describe 'what happened' in the application. An action always has a 'type' property (e.g., 'study/addGoal') and an optional 'payload' containing data needed for the update.")
    
    p = doc.add_paragraph()
    p.add_run("• useSelector & useDispatch Hooks: ").bold = True
    p.add_run("The official bridge between React components and the Redux store.\n"
              "  - useSelector(callback): Subscribes a component to specific parts of the store. When the selected data updates, the component automatically re-renders.\n"
              "  - useDispatch(): Returns the store's dispatch function, allowing components to fire actions (e.g., dispatch(addGoal(data))).")

    add_callout_box(
        doc,
        "Golden Rule of Redux Toolkit: Never mutate state in classic Redux, but IN Redux Toolkit createSlice(), you CAN write mutating syntax (like state.items.push(x) or item.completed = true) because Immer safely converts it to immutable updates. This gives us the best of both worlds: simple readability and strict immutability!",
        title="RTK IMMUTABILITY SUPERPOWER",
        bg_color="EFF6FF",
        border_color="2563EB"
    )

    # --- SECTION 2: ARCHITECTURE & DATA FLOW ---
    doc.add_heading("2. Architecture & Unidirectional Data Flow", level=1)
    doc.add_paragraph(
        "Redux enforces a strict Unidirectional (one-way) Data Flow. This architectural pattern guarantees that data always moves in one predictable loop, making applications exceptionally easy to debug and test."
    )
    
    doc.add_heading("The Step-by-Step Lifecycle of a State Change:", level=2)
    doc.add_paragraph(
        "1. User Event / Trigger: A user interacts with the UI (e.g., clicks the 'Add Goal' button in the DailyTrack dashboard).\n"
        "2. Action Dispatch: The React event handler calls dispatch(addGoal({ title: 'Learn RTK', priority: 'High' })).\n"
        "3. Reducer Processing: The store forwards this action to the studySlice reducer. The reducer matches the action type ('study/addGoal'), takes the payload, and safely pushes the new goal into the state array.\n"
        "4. Store State Update: The central store saves the newly calculated state tree.\n"
        "5. UI Re-rendering: All React components listening to goals via useSelector() are notified of the state change and seamlessly re-render with the latest data."
    )

    doc.add_heading("ASCII Architecture Diagram", level=2)
    diagram_text = (
        "+-------------------------------------------------------------------+\n"
        "|                           REDUX STORE                             |\n"
        "|  +-------------------------------------------------------------+  |\n"
        "|  |  Global State Tree (e.g., study: { goals: [...] }, auth)    |  |\n"
        "|  +-------------------------------------------------------------+  |\n"
        "|           ^                                             |         |\n"
        "|           | (4) New State Saved                         | (5)     |\n"
        "|           |                                             | Select  |\n"
        "|  +--------+------------------+                          | & Read  |\n"
        "|  | Slice Reducers            |                          |         |\n"
        "|  | (Immer mutates proxy)     |                          |         |\n"
        "|  +--------+------------------+                          |         |\n"
        "|           ^                                             v         |\n"
        "|           | (3) Action dispatched             +---------+-------+ |\n"
        "|  +--------+------------------+                |  React UI       | |\n"
        "|  | Action Object             |                |  Components     | |\n"
        "|  | { type, payload }         |                |  (Re-renders)   | |\n"
        "|  +--------+------------------+                +---------+-------+ |\n"
        "|           ^                                             |         |\n"
        "|           | (2) dispatch(action)                        | (1)     |\n"
        "+-----------|---------------------------------------------|---------+\n"
        "            |                                             |\n"
        "            +----- User Interaction (Click / Input) ------+"
    )
    add_code_block(doc, diagram_text)

    # --- SECTION 3: FOLDER STRUCTURE ---
    doc.add_heading("3. Best Practices for Folder Structure", level=1)
    doc.add_paragraph(
        "For modern production applications, the 'Feature-driven' or 'Duck Pattern' folder structure is recommended. Instead of separating files by technical type (e.g., /actions, /reducers), files are grouped by feature domain. Here is the clean folder architecture used in standard RTK projects and our DailyTrack application:"
    )
    
    tree_text = (
        "src/\n"
        "├── components/              # Reusable UI presentation components\n"
        "├── redux/                   # State management root directory\n"
        "│   ├── store.js             # Central store configuration & localStorage persistence\n"
        "│   ├── authSlice.js         # Authentication domain slice (users, login, session)\n"
        "│   ├── studySlice.js        # Study goals domain slice (CRUD, filters, search)\n"
        "│   └── themeSlice.js        # UI theme domain slice (dark/light mode toggle)\n"
        "├── App.jsx                  # Root application layout\n"
        "└── main.jsx                 # Application entry point wrapping <Provider store={store}>"
    )
    add_code_block(doc, tree_text)

    # --- SECTION 4: IMPORTANT FUNCTIONS REFERENCE ---
    doc.add_heading("4. Important RTK Functions & Hooks Reference", level=1)
    doc.add_paragraph("A quick-reference guide to the essential functions and hooks used daily in Redux Toolkit development:")
    
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ["Function / Hook", "Purpose & Practical Usage"]
    for i, title in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    row_data = [
        ("configureStore()", "Wraps standard Redux createStore. Automatically sets up Redux DevTools, combines slice reducers, adds thunk middleware, and enables preloadedState for persistence."),
        ("createSlice()", "The workhorse of RTK. Takes an initial state, slice name, and object of reducer functions. Automatically generates action creators and action types for each reducer."),
        ("createAsyncThunk()", "Handles asynchronous side effects (like fetching data from a REST API). Generates pending, fulfilled, and rejected action lifecycle states automatically."),
        ("useSelector()", "React hook that takes a selector function (e.g., state => state.study.goals) to extract data from the store and subscribe the component to updates."),
        ("useDispatch()", "React hook that returns a reference to the store's dispatch function, enabling components to dispatch actions and trigger state transitions.")
    ]
    
    for row_idx, (fn, desc) in enumerate(row_data, start=1):
        c0 = table.cell(row_idx, 0)
        c1 = table.cell(row_idx, 1)
        
        # Alternating row background
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, top=100, bottom=100, left=120, right=120)
        set_cell_margins(c1, top=100, bottom=100, left=120, right=120)
        
        r0 = c0.paragraphs[0].add_run(fn)
        r0.bold = True
        r0.font.name = 'Consolas'
        r0.font.color.rgb = RGBColor(180, 83, 9) # Brown/Amber
        
        c1.paragraphs[0].add_run(desc)
        
    doc.add_paragraph()

    # --- SECTION 5: REAL-WORLD CASE STUDY: DAILYTRACK ---
    doc.add_heading("5. Real-World Case Study: The DailyTrack Project", level=1)
    doc.add_paragraph(
        "To solidify these concepts, let's examine how Redux Toolkit is implemented in a production-grade project: DailyTrack (dev-track). DailyTrack is a full-featured developer productivity and goal-tracking web application built with React 19, Vite, Tailwind CSS, and Redux Toolkit (@reduxjs/toolkit v2.12.0)."
    )
    
    doc.add_heading("Multi-Slice Architectural Design", level=2)
    doc.add_paragraph(
        "In DailyTrack, global state is divided into three distinct domain slices, preventing state clutter and ensuring separation of concerns:\n\n"
        "1. Study Slice (studySlice.js): Manages the core goal-tracking functionality.\n"
        "   • State Features: An array of goals, filterStatus ('All', 'Completed', 'In Progress'), and searchQuery.\n"
        "   • Reducer Actions: Supports complete CRUD operations including addGoal (using unshift for reverse-chronological ordering), updateGoal, deleteGoal, toggleGoalStatus, and clearCompletedGoals.\n"
        "   • Smart Filtering: Search and status filters are stored in Redux, allowing UI components to dynamically compute visible goals without mutating the master goals list.\n\n"
        "2. Authentication Slice (authSlice.js): Manages user accounts and session state.\n"
        "   • State Features: Stores registered users array and currentUser active session object.\n"
        "   • Validation Logic: Implements custom loginUser and registerUser reducers that check for duplicate email addresses and validate credentials directly within the slice logic, setting descriptive error messages when validation fails.\n\n"
        "3. Theme Slice (themeSlice.js): Manages UI appearance.\n"
        "   • State Features: Stores dark/light theme preference (mode: 'dark' | 'light').\n"
        "   • Actions: Provides toggleTheme and setTheme for instant UI responsiveness."
    )
    
    doc.add_heading("Advanced Technique: State Persistence with LocalStorage", level=2)
    doc.add_paragraph(
        "One of the biggest challenges in SPA development is that Redux store memory resets whenever the browser page is refreshed. DailyTrack solves this elegantly in src/redux/store.js using store subscriptions and preloaded hydration:\n\n"
        "• Hydration on Startup: During configureStore(), the preloadedState property calls loadFromStorage(), which reads serialized JSON strings from localStorage ('devtrack_redux_state_v1') and initializes the store with persistent data.\n"
        "• Automated Subscription: Using store.subscribe(), the application listens to every dispatched action. Whenever state changes, saveToStorage() automatically converts state.study into a JSON string and saves it back to browser storage.\n"
        "• Slice-level Storage: Similarly, authSlice and themeSlice independently synchronize their session and theme keys ('devtrack_users_v1' and 'devtrack_theme_v1') upon action dispatch."
    )

    code_store_example = (
        "// Snippet from DailyTrack src/redux/store.js showing automated persistence\n"
        "export const store = configureStore({\n"
        "  reducer: {\n"
        "    study: studyReducer,\n"
        "    auth: authReducer,\n"
        "    theme: themeReducer,\n"
        "  },\n"
        "  preloadedState: loadFromStorage(), // Hydrate from localStorage on app launch\n"
        "});\n\n"
        "// Subscribe to store updates to save state automatically\n"
        "store.subscribe(() => {\n"
        "  saveToStorage(store.getState());\n"
        "});"
    )
    add_code_block(doc, code_store_example)

    # --- SECTION 6: CHALLENGES & SOLUTIONS ---
    doc.add_heading("6. Challenges Faced & Practical Solutions", level=1)
    doc.add_paragraph("When building scalable apps with Redux Toolkit, developers encounter several practical hurdles. Here is how I overcame them:")
    
    challenges = [
        ("Challenge 1: Mental Shift from Immutable Spreading to Immer Mutability",
         "In early Redux training, we are drilled never to mutate state directly. Switching to RTK's createSlice() felt counter-intuitive at first. The solution was understanding how Immer uses JavaScript Proxies under the hood. Now I confidently write state.goals.unshift(item) knowing RTK guarantees immutability."),
        
        ("Challenge 2: Preventing Data Loss on Browser Page Reloads",
         "Initial prototypes lost all user goals when hitting F5. Instead of littering React components with useEffect localStorage hooks, I implemented the centralized store.subscribe() pattern in store.js (as seen in DailyTrack). This decoupled persistence logic from UI components entirely."),
        
        ("Challenge 3: Action Naming Collisions Across Large Modules",
         "In classic Redux, defining constants like 'ADD_ITEM' in different files caused collisions. RTK solved this automatically: createSlice() prefixes every action creator with the slice name (e.g., 'study/addGoal' vs 'todo/addGoal')."),
        
        ("Challenge 4: Deciding Between Local Component State vs. Global Redux State",
         "A common beginner mistake is putting everything (even form input strings or dropdown open/close booleans) into Redux. The solution is applying the 'Shared State Rule': If only one component cares about it (like a temporary form input), use React's useState. If multiple components or persistent sessions care about it (like user auth, goals list, or theme), put it in Redux.")
    ]
    
    for title, desc in challenges:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {title}\n")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(180, 83, 9)
        p.add_run(desc)

    # --- SECTION 7: ADDITIONAL THINGS EXPLORED ---
    doc.add_heading("7. Additional Things Explored (Beyond Core Basics)", level=1)
    doc.add_paragraph(
        "To become a well-rounded Redux engineer, I explored several advanced features in the Redux ecosystem:\n\n"
        "• Redux Toolkit Query (RTK Query): An advanced data fetching and caching tool built directly into RTK (@reduxjs/toolkit/query). It eliminates the need for writing useEffect data fetching and manual loading/error reducers by providing automated API endpoint definitions, caching, deduplication, and polling.\n"
        "• Redux DevTools Extension Mastery: Learning how to use time-travel debugging, inspect payload diffs across action dispatches, and export/import state traces to reproduce bugs reported by QA teams.\n"
        "• Custom Middleware: Understanding how to intercept actions before they reach reducers (useful for analytics logging, crash reporting, or custom JWT token refreshing)."
    )

    add_callout_box(
        doc,
        "1. Keep store slices focused and domain-specific (auth, study, theme, ui).\n"
        "2. Let RTK handle immutability via Immer—don't write manual {...state} spreads inside createSlice().\n"
        "3. Always wrap root components with <Provider store={store}> in main.jsx.\n"
        "4. Use preloadedState and store.subscribe() for clean, professional state persistence without UI clutter.\n"
        "5. Treat Redux DevTools as your best friend during debugging!",
        title="MY PERSONAL GOLDEN RULES FOR REDUX TOOLKIT",
        bg_color="FEF3C7",
        border_color="D97706"
    )
    
    output_path = os.path.join(os.path.abspath("."), "redux.docx")
    doc.save(output_path)
    print(f"Successfully generated documentation: {output_path}")

if __name__ == "__main__":
    create_redux_documentation()
